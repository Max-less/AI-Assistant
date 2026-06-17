"""
Document loaders for various formats.
Supported: Markdown (.md), PDF (.pdf), DOCX (.docx), TXT (.txt).
"""

import os
import re
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class Document:
    text: str
    source: str  # file path
    metadata: dict = field(default_factory=dict)


def load_markdown(path: str) -> str:
    """Load text from a Markdown file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _clean_pdf_text(text: str) -> str:
    """Tidy raw PDF text: stitch words hyphenated across line breaks and trim
    trailing whitespace, while preserving paragraph structure for the chunker."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # join words split by a hyphen at a line break: "конфигу-\nрацией" -> one word
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+\n", "\n", text)  # trailing spaces
    text = re.sub(r"\n{3,}", "\n\n", text)  # collapse big gaps
    return text.strip()


def load_pdf(path: str) -> str:
    """Extract text from a PDF.

    Uses PyMuPDF, which reconstructs human reading order far better than pypdf —
    pypdf emitted glyphs in content-stream order, which on this corpus produced
    one-word-per-line output and scrambled text on multi-column/diagram pages.
    Falls back to pypdf only if PyMuPDF can't open the file.
    """
    try:
        with fitz.open(path) as doc:
            pages = [page.get_text("text") for page in doc]
        return _clean_pdf_text("\n".join(pages))
    except Exception:
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return _clean_pdf_text("\n".join(pages))


def load_docx(path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def load_txt(path: str) -> str:
    """Load text from a TXT file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


LOADERS = {
    ".md": load_markdown,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
}


def load_document(path: str) -> Document:
    """
    Dispatcher: detect format by file extension and load the document.
    Returns a Document with text, source path, and metadata.
    """
    ext = os.path.splitext(path)[1].lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported format: {ext} ({path})")

    text = loader(path)
    return Document(
        text=text,
        source=path,
        metadata={
            "filename": os.path.basename(path),
            "extension": ext,
            "size_bytes": os.path.getsize(path),
        },
    )